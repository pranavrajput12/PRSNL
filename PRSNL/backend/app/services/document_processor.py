"""Document processing service for handling uploaded files"""
import asyncio
import csv
import hashlib
import io
import logging
import mimetypes
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from uuid import UUID

import aiofiles
import aiofiles.os
import docx

# For text extraction
from markitdown import MarkItDown
import pytesseract
from PIL import Image

# For file type detection
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    magic = None

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing uploaded documents and extracting content"""
    
    def __init__(self, storage_root: str = "storage"):
        self.storage_root = Path(storage_root)
        self.markitdown = MarkItDown()  # Initialize MarkItDown processor
        self.supported_types = {
            'document': ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.csv', '.epub', '.zip'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.webp', '.tiff'],
            'text': ['.txt', '.md', '.rtf', '.xml', '.json'],
            'office': ['.doc', '.docx', '.odt', '.pptx', '.xlsx', '.xls'],
            'pdf': ['.pdf'],
            'spreadsheet': ['.csv', '.xls', '.xlsx'],
            'ebook': ['.epub'],
            'archive': ['.zip']
        }
        
        # Create storage directories
        self._ensure_storage_directories()
    
    async def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from any supported file format using MarkItDown"""
        try:
            logger.info(f"Starting text extraction for: {file_path}")
            
            # Read file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Get file info
            filename = Path(file_path).name
            file_info = self._detect_file_type(file_content, filename)
            
            # Extract content using unified method
            result = await self._extract_content(file_content, file_info, filename)
            
            return {
                "success": True,
                "text": result.get("text", ""),
                "metadata": result.get("metadata", {}),
                "extraction_method": result.get("extraction_method", "unknown"),
                "pages": result.get("pages", 0),
                "word_count": result.get("word_count", 0)
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {}
            }
    
    async def extract_pdf_text(self, file_path: str, use_ocr_fallback: bool = True, 
                              extract_images: bool = False) -> Dict[str, Any]:
        """Extract text specifically from PDF files using MarkItDown with OCR fallback"""
        try:
            logger.info(f"Starting PDF text extraction for: {file_path}")
            
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            filename = Path(file_path).name
            
            # First try MarkItDown
            result = await self._extract_with_markitdown(file_content, '.pdf', filename)
            
            if result and not result.get('error'):
                return {
                    "success": True,
                    "text": result.get("text", ""),
                    "pages": result.get("pages", 0),
                    "metadata": result.get("metadata", {}),
                    "extraction_method": result.get("extraction_method", "markitdown"),
                    "used_ocr": False
                }
            
            # If MarkItDown fails and OCR is enabled, try OCR fallback
            if use_ocr_fallback:
                logger.info(f"MarkItDown failed for {filename}, attempting OCR fallback")
                ocr_result = await self._extract_image_content_fallback(file_content)
                
                if ocr_result and not ocr_result.get('error'):
                    return {
                        "success": True,
                        "text": ocr_result.get("text", ""),
                        "pages": ocr_result.get("pages", 1),
                        "metadata": ocr_result.get("metadata", {}),
                        "extraction_method": "tesseract_ocr",
                        "used_ocr": True
                    }
            
            return {
                "success": False,
                "error": "PDF extraction failed with all methods",
                "text": "",
                "metadata": {}
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {}
            }
    
    async def clean_extracted_text(self, text: str, remove_headers_footers: bool = True) -> str:
        """Clean extracted text by removing headers, footers, and normalizing whitespace"""
        if not text:
            return ""
        
        logger.info(f"Cleaning extracted text ({len(text)} characters)")
        
        cleaned = text
        
        if remove_headers_footers:
            # Remove common header/footer patterns
            lines = cleaned.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line_stripped = line.strip()
                
                # Skip lines that look like headers/footers
                if not (
                    line_stripped.isdigit() or  # Page numbers
                    ('Page' in line_stripped and any(char.isdigit() for char in line_stripped)) or
                    len(line_stripped) < 3 or  # Very short lines
                    line_stripped.startswith('___') or  # Underline separators
                    line_stripped.startswith('---') or  # Dash separators
                    line_stripped.count('|') > 3  # Table separators
                ):
                    cleaned_lines.append(line)
            
            cleaned = '\n'.join(cleaned_lines)
        
        # Normalize whitespace while preserving paragraph breaks
        paragraphs = cleaned.split('\n\n')
        normalized_paragraphs = []
        
        for paragraph in paragraphs:
            # Normalize whitespace within paragraphs
            normalized = ' '.join(paragraph.split())
            if normalized:  # Only add non-empty paragraphs
                normalized_paragraphs.append(normalized)
        
        cleaned = '\n\n'.join(normalized_paragraphs)
        
        logger.info(f"Text cleaning completed ({len(cleaned)} characters after cleaning)")
        
        return cleaned
    
    def _ensure_storage_directories(self):
        """Create storage directory structure"""
        directories = [
            self.storage_root / "documents",
            self.storage_root / "images", 
            self.storage_root / "extracted_text",
            self.storage_root / "thumbnails",
            self.storage_root / "temp"
        ]
        
        for directory in directories:
            directory.mkdir(parents=True, exist_ok=True)
    
    def _get_file_hash(self, file_content: bytes) -> str:
        """Generate SHA-256 hash for file content"""
        return hashlib.sha256(file_content).hexdigest()
    
    def _detect_file_type(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Detect file type using multiple methods"""
        # Get file extension
        file_extension = Path(filename).suffix.lower()
        
        # Use python-magic for MIME type detection if available
        if HAS_MAGIC:
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
            except Exception as e:
                logger.warning(f"Failed to detect MIME type with magic: {e}")
                mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        else:
            # Fallback to mimetypes module
            mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        
        # Categorize file type
        category = self._categorize_file_type(file_extension, mime_type)
        
        return {
            'extension': file_extension,
            'mime_type': mime_type,
            'category': category,
            'size': len(file_content)
        }
    
    def _categorize_file_type(self, extension: str, mime_type: str) -> str:
        """Categorize file based on extension and MIME type"""
        # All document-like files should be categorized as 'document'
        if extension in ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.csv', '.epub', '.zip', '.xml', '.json']:
            return 'document'
        elif extension in ['.pptx', '.xlsx', '.xls']:
            return 'office'
        elif extension in self.supported_types['image']:
            return 'image'
        elif extension in self.supported_types['document']:
            return 'document'
        else:
            return 'unknown'
    
    def _generate_file_path(self, item_id: UUID, file_hash: str, extension: str, category: str) -> Path:
        """Generate organized file path"""
        # Create date-based subdirectory
        date_dir = datetime.now().strftime("%Y/%m")
        
        # Base directory based on category
        if category == 'image':
            base_dir = self.storage_root / "images" / date_dir
        else:
            base_dir = self.storage_root / "documents" / date_dir
        
        base_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with hash to prevent duplicates
        filename = f"{item_id}_{file_hash[:8]}{extension}"
        return base_dir / filename
    
    async def process_file(self, file_content: bytes, filename: str, item_id: UUID) -> Dict[str, Any]:
        """Main file processing pipeline"""
        logger.info(f"Processing file: {filename} for item {item_id}")
        
        # Detect file type
        file_info = self._detect_file_type(file_content, filename)
        file_hash = self._get_file_hash(file_content)
        
        # Generate storage path
        file_path = self._generate_file_path(
            item_id, file_hash, file_info['extension'], file_info['category']
        )
        
        # Save file to storage
        await self._save_file(file_content, file_path)
        
        # Extract content based on file type
        extracted_content = await self._extract_content(file_content, file_info, filename)
        
        # Generate thumbnail for images
        thumbnail_path = None
        if file_info['category'] == 'image':
            thumbnail_path = await self._generate_thumbnail(file_content, item_id, file_hash)
        
        # Save extracted text
        text_file_path = None
        if extracted_content.get('text'):
            text_file_path = await self._save_extracted_text(
                extracted_content['text'], item_id, file_hash
            )
        
        return {
            'file_path': str(file_path),
            'file_hash': file_hash,
            'file_info': file_info,
            'extracted_content': extracted_content,
            'thumbnail_path': str(thumbnail_path) if thumbnail_path else None,
            'text_file_path': str(text_file_path) if text_file_path else None,
            'original_filename': filename,
            'processed_at': datetime.now().isoformat()
        }
    
    async def _save_file(self, file_content: bytes, file_path: Path):
        """Save file to storage"""
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        logger.info(f"Saved file to: {file_path}")
    
    async def _extract_content(self, file_content: bytes, file_info: Dict, filename: str) -> Dict[str, Any]:
        """Extract content from file using MarkItDown as primary processor with intelligent fallbacks"""
        extension = file_info['extension']
        
        # First, try MarkItDown for all supported formats
        markitdown_result = await self._extract_with_markitdown(file_content, extension, filename)
        if markitdown_result and not markitdown_result.get('error'):
            return markitdown_result
        
        # If MarkItDown fails, use intelligent fallbacks
        logger.warning(f"MarkItDown failed for {filename}, using fallback extraction")
        
        try:
            category = file_info['category']
            
            # Fallback extraction methods
            if category == 'document':
                if extension == '.csv':
                    return await self._extract_csv_content(file_content)
                elif extension in ['.txt', '.rtf', '.xml', '.json']:
                    return await self._extract_text_content(file_content)
                elif extension in ['.doc', '.docx', '.odt']:
                    return await self._extract_office_content_fallback(file_content, extension)
                else:
                    return await self._extract_text_content(file_content)
            elif category == 'office':
                return await self._extract_office_content_fallback(file_content, extension)
            elif category == 'image':
                return await self._extract_image_content_fallback(file_content)
            else:
                logger.warning(f"Unsupported file type: {category}")
                return {
                    'text': f"Unsupported file type: {filename}",
                    'pages': 0,
                    'word_count': 0,
                    'extraction_method': 'unsupported'
                }
        except Exception as e:
            logger.error(f"Content extraction failed for {filename}: {e}")
            return {
                'text': f"Failed to extract content from: {filename}",
                'pages': 0,
                'word_count': 0,
                'extraction_method': 'failed',
                'error': str(e)
            }
    
    async def _extract_with_markitdown(self, file_content: bytes, extension: str, filename: str) -> Dict[str, Any]:
        """Extract content using MarkItDown for all supported formats"""
        try:
            # Create temporary file for MarkItDown processing
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Process with MarkItDown
                result = self.markitdown.convert(temp_file_path)
                
                # Extract text content
                text_content = result.text_content if hasattr(result, 'text_content') else str(result)
                
                # Determine pages count based on file type
                pages = self._estimate_pages(text_content, extension)
                
                # Create comprehensive metadata
                metadata = {
                    'source_type': self._get_source_type(extension),
                    'content_type': getattr(result, 'content_type', 'unknown'),
                    'title': getattr(result, 'title', ''),
                    'preserves_formatting': True,
                    'unified_processor': True,
                    'markitdown_version': 'latest'
                }
                
                # Add format-specific metadata
                if extension in ['.epub']:
                    metadata['format'] = 'ebook'
                elif extension in ['.zip']:
                    metadata['format'] = 'archive'
                elif extension in ['.pptx']:
                    metadata['format'] = 'presentation'
                elif extension in ['.xlsx', '.xls']:
                    metadata['format'] = 'spreadsheet'
                
                return {
                    'text': text_content,
                    'pages': pages,
                    'word_count': len(text_content.split()) if text_content else 0,
                    'extraction_method': 'markitdown_unified',
                    'metadata': metadata
                }
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass
                    
        except Exception as e:
            logger.error(f"MarkItDown extraction failed for {filename}: {e}")
            return {
                'error': f"MarkItDown extraction failed: {str(e)}",
                'extraction_method': 'markitdown_failed'
            }
    
    def _estimate_pages(self, text_content: str, extension: str) -> int:
        """Estimate page count based on content and file type"""
        if not text_content:
            return 0
        
        # Different estimation strategies based on file type
        if extension == '.pdf':
            return max(1, len(text_content) // 2000)  # ~2000 chars per page
        elif extension in ['.pptx']:
            # PowerPoint slides - estimate based on content sections
            return max(1, text_content.count('\n\n') + 1)
        elif extension in ['.xlsx', '.xls']:
            # Spreadsheet sheets - estimate based on table structures
            return max(1, text_content.count('---') + 1)
        elif extension in ['.epub']:
            # E-book chapters
            return max(1, len(text_content) // 3000)  # ~3000 chars per e-book page
        elif extension in ['.zip']:
            # Archive files - count of contained files
            return max(1, text_content.count('File:'))
        else:
            return 1  # Single page for most other formats
    
    def _get_source_type(self, extension: str) -> str:
        """Get source type description for metadata"""
        type_map = {
            '.pdf': 'pdf_document',
            '.docx': 'word_document', 
            '.doc': 'word_document',
            '.pptx': 'powerpoint_presentation',
            '.xlsx': 'excel_spreadsheet',
            '.xls': 'excel_spreadsheet',
            '.epub': 'ebook',
            '.zip': 'archive',
            '.xml': 'structured_data',
            '.json': 'structured_data',
            '.txt': 'plain_text',
            '.rtf': 'rich_text',
            '.csv': 'tabular_data'
        }
        return type_map.get(extension, 'unknown')

    
    async def _extract_office_content_fallback(self, file_content: bytes, extension: str) -> Dict[str, Any]:
        """Fallback extraction for Office documents using python-docx"""
        try:
            # Only handle .docx files with python-docx fallback
            if extension == '.docx':
                doc = docx.Document(io.BytesIO(file_content))
                text_content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                full_text = '\n'.join(text_content)
                
                return {
                    'text': full_text,
                    'pages': len(doc.paragraphs),
                    'word_count': len(full_text.split()) if full_text else 0,
                    'extraction_method': 'python-docx_fallback',
                    'metadata': {
                        'paragraphs': len(doc.paragraphs),
                        'tables': len(doc.tables),
                        'fallback_reason': 'markitdown_failed'
                    }
                }
            else:
                # For other office formats, return error since we don't have fallbacks
                return {
                    'text': f"No fallback available for {extension} files",
                    'pages': 1,
                    'word_count': 0,
                    'extraction_method': 'no_fallback',
                    'error': f"Unsupported format for fallback: {extension}"
                }
                
        except Exception as e:
            logger.error(f"Office document fallback extraction failed: {e}")
            return {
                'text': f"Failed to extract office content: {str(e)}",
                'pages': 1,
                'word_count': 0,
                'extraction_method': 'fallback_failed',
                'error': str(e)
            }
    
    async def _extract_text_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from plain text files with enhanced structured data handling"""
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                text = file_content.decode('latin-1')
            except UnicodeDecodeError:
                text = file_content.decode('utf-8', errors='replace')
        
        # Enhanced processing for structured data
        extraction_method = 'text_decode'
        metadata = {
            'encoding': 'utf-8',
            'characters': len(text)
        }
        
        # Check if it's JSON or XML and format it nicely
        try:
            import json
            if text.strip().startswith('{') or text.strip().startswith('['):
                # Try to parse and pretty-print JSON
                parsed_json = json.loads(text)
                text = json.dumps(parsed_json, indent=2)
                extraction_method = 'json_formatted'
                metadata['data_type'] = 'json'
                metadata['json_valid'] = True
        except (json.JSONDecodeError, ValueError):
            pass
        
        try:
            import xml.etree.ElementTree as ET
            if text.strip().startswith('<'):
                # Try to parse and format XML
                root = ET.fromstring(text)
                # Convert to a more readable format
                text = f"XML Document: {root.tag}\n" + self._xml_to_text(root)
                extraction_method = 'xml_formatted'
                metadata['data_type'] = 'xml'
                metadata['xml_valid'] = True
        except ET.ParseError:
            pass
        
        return {
            'text': text,
            'pages': 1,
            'word_count': len(text.split()) if text else 0,
            'extraction_method': extraction_method,
            'metadata': metadata
        }
    
    def _xml_to_text(self, element, level=0) -> str:
        """Convert XML element to readable text format"""
        result = []
        indent = "  " * level
        
        if element.text and element.text.strip():
            result.append(f"{indent}{element.tag}: {element.text.strip()}")
        else:
            result.append(f"{indent}{element.tag}:")
        
        for child in element:
            result.append(self._xml_to_text(child, level + 1))
        
        return "\n".join(result)
    
    async def _extract_csv_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract content from CSV files"""
        try:
            # Decode CSV content
            text = file_content.decode('utf-8')
            
            # Parse CSV to create a readable text representation
            csv_reader = csv.reader(io.StringIO(text))
            rows = list(csv_reader)
            
            # Create a text representation of the CSV
            text_lines = []
            if rows:
                # Add headers if present
                headers = rows[0]
                text_lines.append("CSV Data with columns: " + ", ".join(headers))
                text_lines.append("")
                
                # Add data rows
                for row in rows[1:]:
                    text_lines.append(" | ".join(row))
            
            extracted_text = "\n".join(text_lines)
            
            return {
                'text': extracted_text,
                'pages': 1,
                'word_count': len(extracted_text.split()) if extracted_text else 0,
                'extraction_method': 'csv_parse',
                'metadata': {
                    'row_count': len(rows),
                    'column_count': len(rows[0]) if rows else 0,
                    'encoding': 'utf-8'
                }
            }
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            # Fallback to plain text extraction
            return await self._extract_text_content(file_content)
    
    async def _extract_image_content_fallback(self, file_content: bytes) -> Dict[str, Any]:
        """Fallback extraction for images using Tesseract OCR"""
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            extracted_text = pytesseract.image_to_string(image)
            
            return {
                'text': extracted_text,
                'pages': 1,
                'word_count': len(extracted_text.split()) if extracted_text else 0,
                'extraction_method': 'tesseract_ocr_fallback',
                'metadata': {
                    'image_size': image.size,
                    'image_mode': image.mode,
                    'has_text': bool(extracted_text.strip()),
                    'fallback_reason': 'markitdown_failed'
                }
            }
        except Exception as e:
            logger.error(f"Tesseract OCR fallback failed: {e}")
            return {
                'text': "Image content (OCR not available)",
                'pages': 1,
                'word_count': 0,
                'extraction_method': 'ocr_failed',
                'error': str(e)
            }
    
    def _detect_image_format(self, file_content: bytes) -> str:
        """Detect image format from file content"""
        try:
            image = Image.open(io.BytesIO(file_content))
            format_map = {
                'JPEG': '.jpg',
                'PNG': '.png',
                'GIF': '.gif',
                'BMP': '.bmp',
                'TIFF': '.tiff',
                'WEBP': '.webp'
            }
            return format_map.get(image.format, '.jpg')
        except Exception:
            return '.jpg'  # Default fallback
    
    async def _generate_thumbnail(self, file_content: bytes, item_id: UUID, file_hash: str) -> Optional[Path]:
        """Generate thumbnail for image files"""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Create thumbnail
            thumbnail_size = (300, 300)
            image.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
            
            # Save thumbnail
            thumbnail_dir = self.storage_root / "thumbnails" / datetime.now().strftime("%Y/%m")
            thumbnail_dir.mkdir(parents=True, exist_ok=True)
            
            thumbnail_path = thumbnail_dir / f"{item_id}_{file_hash[:8]}_thumb.jpg"
            image.save(thumbnail_path, 'JPEG', quality=85)
            
            return thumbnail_path
        except Exception as e:
            logger.error(f"Thumbnail generation failed: {e}")
            return None
    
    async def _save_extracted_text(self, text: str, item_id: UUID, file_hash: str) -> Path:
        """Save extracted text to file"""
        text_dir = self.storage_root / "extracted_text" / datetime.now().strftime("%Y/%m")
        text_dir.mkdir(parents=True, exist_ok=True)
        
        text_file_path = text_dir / f"{item_id}_{file_hash[:8]}.txt"
        
        async with aiofiles.open(text_file_path, 'w', encoding='utf-8') as f:
            await f.write(text)
        
        return text_file_path
    
    async def get_file_content(self, file_path: str) -> Optional[bytes]:
        """Retrieve file content from storage"""
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                return await f.read()
        except FileNotFoundError:
            logger.error(f"File not found: {file_path}")
            return None
        except Exception as e:
            logger.error(f"Failed to read file {file_path}: {e}")
            return None
    
    async def delete_file(self, file_path: str) -> bool:
        """Delete file from storage"""
        try:
            await aiofiles.os.remove(file_path)
            logger.info(f"Deleted file: {file_path}")
            return True
        except FileNotFoundError:
            logger.warning(f"File not found for deletion: {file_path}")
            return False
        except Exception as e:
            logger.error(f"Failed to delete file {file_path}: {e}")
            return False
    
    def get_storage_stats(self) -> Dict[str, Any]:
        """Get storage usage statistics"""
        stats = {
            'total_files': 0,
            'total_size_bytes': 0,
            'by_category': {}
        }
        
        for category_dir in ['documents', 'images', 'thumbnails', 'extracted_text']:
            category_path = self.storage_root / category_dir
            if category_path.exists():
                category_stats = self._get_directory_stats(category_path)
                stats['by_category'][category_dir] = category_stats
                stats['total_files'] += category_stats['files']
                stats['total_size_bytes'] += category_stats['size_bytes']
        
        return stats
    
    def _get_directory_stats(self, directory: Path) -> Dict[str, Any]:
        """Get statistics for a directory"""
        total_size = 0
        file_count = 0
        
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                file_count += 1
                total_size += file_path.stat().st_size
        
        return {
            'files': file_count,
            'size_bytes': total_size,
            'size_mb': round(total_size / (1024 * 1024), 2)
        }